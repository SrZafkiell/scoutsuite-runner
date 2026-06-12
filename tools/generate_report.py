import argparse
import csv
import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from docx import Document


SEVERITY_MAP = {
    "danger": "High",
    "warning": "Medium",
    "info": "Informational",
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
    "informational": "Informational",
}

SEVERITY_ORDER = {
    "Critical": 0,
    "High": 1,
    "Medium": 2,
    "Low": 3,
    "Informational": 4,
    "Unknown": 5,
}


def clean_text(value: Any) -> str:
    if value is None:
        return ""

    text = str(value)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def md_escape(value: Any) -> str:
    text = clean_text(value)
    return text.replace("|", "\\|")


def find_scoutsuite_results_file(run_dir: Path) -> Path:
    candidates = list(run_dir.glob("scoutsuite-results/scoutsuite_results_*.js"))

    if not candidates:
        candidates = list(run_dir.glob("**/scoutsuite_results_*.js"))

    if not candidates:
        raise FileNotFoundError(
            f"No ScoutSuite results file found under: {run_dir}"
        )

    return sorted(candidates, key=lambda path: path.stat().st_mtime, reverse=True)[0]


def load_scoutsuite_js(path: Path) -> dict[str, Any]:
    raw = path.read_text(encoding="utf-8", errors="replace")

    start = raw.find("{")
    end = raw.rfind("}")

    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"Could not extract JSON object from: {path}")

    json_text = raw[start : end + 1]
    return json.loads(json_text)


def walk_findings(node: Any, path: list[str] | None = None):
    path = path or []

    if isinstance(node, dict):
        findings = node.get("findings")

        if isinstance(findings, dict):
            for finding_id, finding in findings.items():
                if isinstance(finding, dict):
                    yield path + ["findings", str(finding_id)], str(finding_id), finding

        for key, value in node.items():
            yield from walk_findings(value, path + [str(key)])

    elif isinstance(node, list):
        for index, value in enumerate(node):
            yield from walk_findings(value, path + [str(index)])


def get_service_from_path(path: list[str], finding: dict[str, Any]) -> str:
    explicit = finding.get("service") or finding.get("service_name")
    if explicit:
        return clean_text(explicit)

    if "services" in path:
        index = path.index("services")
        if len(path) > index + 1:
            return clean_text(path[index + 1]).upper()

    return "Unknown"


def get_severity(finding: dict[str, Any]) -> str:
    raw = (
        finding.get("level")
        or finding.get("severity")
        or finding.get("risk_level")
        or finding.get("risk")
        or "Unknown"
    )

    normalized = clean_text(raw).lower()
    return SEVERITY_MAP.get(normalized, clean_text(raw).title() or "Unknown")


def get_title(finding_id: str, finding: dict[str, Any]) -> str:
    return clean_text(
        finding.get("dashboard_name")
        or finding.get("display_name")
        or finding.get("title")
        or finding.get("description")
        or finding_id.replace("_", " ").replace("-", " ").title()
    )


def get_description(finding: dict[str, Any]) -> str:
    return clean_text(
        finding.get("description")
        or finding.get("rationale")
        or finding.get("rationale_description")
        or ""
    )


def get_recommendation(finding: dict[str, Any]) -> str:
    return clean_text(
        finding.get("remediation")
        or finding.get("recommendation")
        or finding.get("fix")
        or finding.get("references")
        or "Review the affected resource and apply the relevant AWS security best practice."
    )


def extract_resource_name(item: Any) -> str:
    if isinstance(item, str):
        return item

    if isinstance(item, dict):
        for key in [
            "arn",
            "id",
            "name",
            "Name",
            "resource_id",
            "resource_name",
            "instance_id",
            "group_id",
            "bucket_name",
            "db_instance_identifier",
            "function_name",
        ]:
            if key in item and item[key]:
                return clean_text(item[key])

        compact = json.dumps(item, ensure_ascii=False, default=str)
        return compact[:160]

    return clean_text(item)


def get_items(finding: dict[str, Any]) -> list[Any]:
    for key in ["items", "flagged_items", "affected_resources", "resources"]:
        value = finding.get(key)

        if isinstance(value, list):
            return value

        if isinstance(value, dict):
            return list(value.keys())

    return []


def build_findings(
    data: dict[str, Any],
    project_filter: str = "",
    min_severity: str = "Informational",
) -> list[dict[str, Any]]:
    findings = []
    seen = set()

    min_severity_rank = SEVERITY_ORDER.get(min_severity, SEVERITY_ORDER["Informational"])

    for path, finding_id, finding in walk_findings(data):
        service = get_service_from_path(path, finding)
        severity = get_severity(finding)

        severity_rank = SEVERITY_ORDER.get(severity, SEVERITY_ORDER["Unknown"])

        if severity_rank > min_severity_rank:
            continue

        title = get_title(finding_id, finding)
        description = get_description(finding)
        recommendation = get_recommendation(finding)

        items = get_items(finding)
        resources = [extract_resource_name(item) for item in items]
        resources = [resource for resource in resources if resource]

        affected_count = len(resources)

        key = (
            service,
            finding_id,
            severity,
            title,
            "|".join(resources[:10]),
        )

        if key in seen:
            continue

        seen.add(key)

        record = {
            "severity": severity,
            "service": service,
            "finding_id": finding_id,
            "title": title,
            "affected_count": affected_count,
            "affected_resources": ", ".join(resources[:25]),
            "description": description,
            "recommendation": recommendation,
            "path": ".".join(path),
            "raw": finding,
        }

        if project_filter:
            needle = project_filter.lower()
            haystack = json.dumps(record, ensure_ascii=False, default=str).lower()

            if needle not in haystack:
                continue

        findings.append(record)

    findings.sort(
        key=lambda item: (
            SEVERITY_ORDER.get(item["severity"], 99),
            item["service"],
            item["title"],
        )
    )

    return findings


def priority_from_severity(severity: str) -> str:
    if severity in ["Critical", "High"]:
        return "P1"
    if severity == "Medium":
        return "P2"
    return "P3"


def write_csv(path: Path, findings: list[dict[str, Any]]) -> None:
    fields = [
        "severity",
        "service",
        "finding_id",
        "title",
        "affected_count",
        "affected_resources",
        "description",
        "recommendation",
        "path",
    ]

    with path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=fields)
        writer.writeheader()

        for finding in findings:
            writer.writerow({field: finding.get(field, "") for field in fields})


def write_tasks_md(path: Path, client: str, findings: list[dict[str, Any]]) -> None:
    grouped = defaultdict(list)

    for finding in findings:
        grouped[priority_from_severity(finding["severity"])].append(finding)

    lines = [
        f"# Remediation Tasks - {client}",
        "",
        "> Draft generated from ScoutSuite output. Validate scope and affected resources before assigning work.",
        "",
    ]

    for priority in ["P1", "P2", "P3"]:
        items = grouped.get(priority, [])

        lines.append(f"## {priority}")
        lines.append("")

        if not items:
            lines.append("No findings in this priority.")
            lines.append("")
            continue

        for finding in items:
            lines.extend(
                [
                    f"- [ ] **{md_escape(finding['title'])}**",
                    f"  - Severity: {md_escape(finding['severity'])}",
                    f"  - Service: {md_escape(finding['service'])}",
                    f"  - Affected resources: {md_escape(finding['affected_resources'] or 'Pending validation')}",
                    f"  - Recommended action: {md_escape(finding['recommendation'])}",
                    f"  - Status: Pending validation",
                    "",
                ]
            )

    path.write_text("\n".join(lines), encoding="utf-8")


def write_report_md(
    path: Path,
    client: str,
    run_dir: Path,
    source_file: Path,
    findings: list[dict[str, Any]],
    project_filter: str,
) -> None:
    counts = Counter(finding["severity"] for finding in findings)
    services = Counter(finding["service"] for finding in findings)

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        f"# Meta Data Security Assessment Draft - {client}",
        "",
        f"Generated at: {generated_at}",
        "",
        "## 1. Executive Summary",
        "",
        (
            "This document is a draft security assessment generated from ScoutSuite output. "
            "It must be manually reviewed before being shared externally."
        ),
        "",
        "## 2. Scope and Boundaries",
        "",
        (
            f"The assessment focuses on AWS resources related to **{client}** and shared infrastructure "
            "that directly affects the application's security posture."
        ),
        "",
        (
            "ScoutSuite was executed against the AWS account using read-only audit credentials. "
            "Because the account may host multiple products, this draft should be filtered and validated "
            "against the actual application inventory before final delivery."
        ),
        "",
    ]

    if project_filter:
        lines.extend(
            [
                f"Project filter used: `{project_filter}`",
                "",
                (
                    "> Warning: name/tag filtering can miss shared resources such as VPCs, security groups, "
                    "IAM policies, KMS keys, logs, or DNS records that affect the application."
                ),
                "",
            ]
        )

    lines.extend(
        [
            "## 3. Source Evidence",
            "",
            f"- ScoutSuite run folder: `{run_dir}`",
            f"- Parsed results file: `{source_file}`",
            "- Validation status: Pending manual review",
            "",
            "## 4. Findings Summary",
            "",
            "| Severity | Count |",
            "|---|---:|",
        ]
    )

    for severity in ["Critical", "High", "Medium", "Low", "Informational", "Unknown"]:
        lines.append(f"| {severity} | {counts.get(severity, 0)} |")

    lines.extend(
        [
            "",
            "## 5. Findings by Service",
            "",
            "| Service | Count |",
            "|---|---:|",
        ]
    )

    for service, count in sorted(services.items()):
        lines.append(f"| {md_escape(service)} | {count} |")

    lines.extend(
        [
            "",
            "## 6. Meta Data Security Control Review",
            "",
            "Use this section to map infrastructure evidence to Meta data security expectations.",
            "",
            "| Control Area | Status | Evidence / Notes |",
            "|---|---|---|",
            "| Access management | Pending review | IAM users, roles, policies, MFA, least privilege |",
            "| Encryption in transit | Pending review | HTTPS, ACM, ALB/CloudFront TLS configuration |",
            "| Encryption at rest | Pending review | EBS, RDS, S3, KMS configuration |",
            "| Credential and token protection | Pending review | Secrets Manager, SSM, environment variables, key rotation |",
            "| Logging and monitoring | Pending review | CloudTrail, CloudWatch, ALB logs, application logs |",
            "| Vulnerability and patch management | Pending review | OS updates, dependency updates, runtime versions |",
            "| Incident response | Pending review | Alerting, escalation, response ownership |",
            "",
            "## 7. Prioritized Remediation Plan",
            "",
            "| Priority | Severity | Service | Finding | Affected Resources |",
            "|---|---|---|---|---|",
        ]
    )

    for finding in findings:
        lines.append(
            "| "
            + " | ".join(
                [
                    priority_from_severity(finding["severity"]),
                    md_escape(finding["severity"]),
                    md_escape(finding["service"]),
                    md_escape(finding["title"]),
                    md_escape(finding["affected_resources"] or "Pending validation"),
                ]
            )
            + " |"
        )

    lines.extend(
        [
            "",
            "## 8. Detailed Findings",
            "",
        ]
    )

    for index, finding in enumerate(findings, start=1):
        lines.extend(
            [
                f"### {index}. {finding['title']}",
                "",
                f"- Severity: {finding['severity']}",
                f"- Priority: {priority_from_severity(finding['severity'])}",
                f"- Service: {finding['service']}",
                f"- Finding ID: `{finding['finding_id']}`",
                f"- Affected resources: {finding['affected_resources'] or 'Pending validation'}",
                "",
                "**Description**",
                "",
                finding["description"] or "No description available in parsed ScoutSuite output.",
                "",
                "**Recommended Action**",
                "",
                finding["recommendation"],
                "",
                "**Validation Notes**",
                "",
                "- [ ] Confirm this resource belongs to the application scope.",
                "- [ ] Confirm the finding is still active.",
                "- [ ] Confirm remediation owner and target date.",
                "",
            ]
        )

    lines.extend(
        [
            "## 9. Appendix",
            "",
            "- Raw ScoutSuite HTML report should be attached or referenced separately.",
            "- Generated CSV and remediation task files are available in the `generated` folder.",
            "",
        ]
    )

    path.write_text("\n".join(lines), encoding="utf-8")


def write_report_docx(
    path: Path,
    client: str,
    run_dir: Path,
    source_file: Path,
    findings: list[dict[str, Any]],
    project_filter: str,
    min_severity: str,
) -> None:
    counts = Counter(finding["severity"] for finding in findings)
    services = Counter(finding["service"] for finding in findings)
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()

    document.add_heading(f"Meta Data Security Assessment Draft - {client}", 0)

    document.add_paragraph(f"Generated at: {generated_at}")
    document.add_paragraph(f"Minimum severity: {min_severity}")
    document.add_paragraph("Validation status: Pending manual review")

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "This document is a draft security assessment generated from ScoutSuite output. "
        "It must be manually reviewed before being shared externally."
    )

    document.add_heading("2. Scope and Boundaries", level=1)
    document.add_paragraph(
        f"The assessment focuses on AWS resources related to {client} and shared infrastructure "
        "that directly affects the application's security posture."
    )
    document.add_paragraph(
        "ScoutSuite was executed against the AWS account using read-only audit credentials. "
        "Because the account may host multiple products, this draft should be filtered and validated "
        "against the actual application inventory before final delivery."
    )

    if project_filter:
        document.add_paragraph(f"Project filter used: {project_filter}")
        document.add_paragraph(
            "Warning: name/tag filtering can miss shared resources such as VPCs, security groups, "
            "IAM policies, KMS keys, logs, or DNS records that affect the application."
        )

    document.add_heading("3. Source Evidence", level=1)
    document.add_paragraph(f"ScoutSuite run folder: {run_dir}")
    document.add_paragraph(f"Parsed results file: {source_file}")

    document.add_heading("4. Findings Summary", level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Severity"
    header_cells[1].text = "Count"

    for severity in ["Critical", "High", "Medium", "Low", "Informational", "Unknown"]:
        row_cells = table.add_row().cells
        row_cells[0].text = severity
        row_cells[1].text = str(counts.get(severity, 0))

    document.add_heading("5. Findings by Service", level=1)

    service_table = document.add_table(rows=1, cols=2)
    service_table.style = "Table Grid"
    header_cells = service_table.rows[0].cells
    header_cells[0].text = "Service"
    header_cells[1].text = "Count"

    for service, count in sorted(services.items()):
        row_cells = service_table.add_row().cells
        row_cells[0].text = service
        row_cells[1].text = str(count)

    document.add_heading("6. Meta Data Security Control Review", level=1)

    control_table = document.add_table(rows=1, cols=3)
    control_table.style = "Table Grid"
    header_cells = control_table.rows[0].cells
    header_cells[0].text = "Control Area"
    header_cells[1].text = "Status"
    header_cells[2].text = "Evidence / Notes"

    controls = [
        ("Access management", "Pending review", "IAM users, roles, policies, MFA, least privilege"),
        ("Encryption in transit", "Pending review", "HTTPS, ACM, ALB/CloudFront TLS configuration"),
        ("Encryption at rest", "Pending review", "EBS, RDS, S3, KMS configuration"),
        ("Credential and token protection", "Pending review", "Secrets Manager, SSM, environment variables, key rotation"),
        ("Logging and monitoring", "Pending review", "CloudTrail, CloudWatch, ALB logs, application logs"),
        ("Vulnerability and patch management", "Pending review", "OS updates, dependency updates, runtime versions"),
        ("Incident response", "Pending review", "Alerting, escalation, response ownership"),
    ]

    for control, status, notes in controls:
        row_cells = control_table.add_row().cells
        row_cells[0].text = control
        row_cells[1].text = status
        row_cells[2].text = notes

    document.add_heading("7. Prioritized Remediation Plan", level=1)

    remediation_table = document.add_table(rows=1, cols=5)
    remediation_table.style = "Table Grid"
    header_cells = remediation_table.rows[0].cells
    header_cells[0].text = "Priority"
    header_cells[1].text = "Severity"
    header_cells[2].text = "Service"
    header_cells[3].text = "Finding"
    header_cells[4].text = "Affected Resources"

    for finding in findings:
        row_cells = remediation_table.add_row().cells
        row_cells[0].text = priority_from_severity(finding["severity"])
        row_cells[1].text = finding["severity"]
        row_cells[2].text = finding["service"]
        row_cells[3].text = finding["title"]
        row_cells[4].text = finding["affected_resources"] or "Pending validation"

    document.add_heading("8. Detailed Findings", level=1)

    for index, finding in enumerate(findings, start=1):
        document.add_heading(f"{index}. {finding['title']}", level=2)

        document.add_paragraph(f"Severity: {finding['severity']}")
        document.add_paragraph(f"Priority: {priority_from_severity(finding['severity'])}")
        document.add_paragraph(f"Service: {finding['service']}")
        document.add_paragraph(f"Finding ID: {finding['finding_id']}")
        document.add_paragraph(f"Affected resources: {finding['affected_resources'] or 'Pending validation'}")

        document.add_paragraph("Description", style="Heading 3")
        document.add_paragraph(
            finding["description"] or "No description available in parsed ScoutSuite output."
        )

        document.add_paragraph("Recommended Action", style="Heading 3")
        document.add_paragraph(finding["recommendation"])

        document.add_paragraph("Validation Notes", style="Heading 3")
        document.add_paragraph("- Confirm this resource belongs to the application scope.")
        document.add_paragraph("- Confirm the finding is still active.")
        document.add_paragraph("- Confirm remediation owner and target date.")

    document.add_heading("9. Appendix", level=1)
    document.add_paragraph("Raw ScoutSuite HTML report should be attached or referenced separately.")
    document.add_paragraph("Generated CSV and remediation task files are available in the generated folder.")

    document.save(path)

def write_summary_json(path: Path, findings: list[dict[str, Any]]) -> None:
    safe_findings = []

    for finding in findings:
        copied = dict(finding)
        copied.pop("raw", None)
        safe_findings.append(copied)

    summary = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "total_findings": len(findings),
        "severity_counts": dict(Counter(finding["severity"] for finding in findings)),
        "service_counts": dict(Counter(finding["service"] for finding in findings)),
        "findings": safe_findings,
    }

    path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate draft report files from ScoutSuite results."
    )

    parser.add_argument("--run-dir", required=True, help="ScoutSuite run directory.")
    parser.add_argument("--client", required=True, help="Client or project name.")
    parser.add_argument(
        "--project-filter",
        default="",
        help="Optional text filter for project/resource names. Use carefully.",
    )
    parser.add_argument(
        "--min-severity",
        default="Informational",
        choices=["Critical", "High", "Medium", "Low", "Informational"],
        help="Minimum severity to include in generated outputs.",
    )

    args = parser.parse_args()

    run_dir = Path(args.run_dir).resolve()
    source_file = find_scoutsuite_results_file(run_dir)
    data = load_scoutsuite_js(source_file)

    findings = build_findings(
        data,
        project_filter=args.project_filter,
        min_severity=args.min_severity,
    )

    generated_dir = run_dir / "generated"
    generated_dir.mkdir(parents=True, exist_ok=True)

    write_csv(generated_dir / "findings.csv", findings)
    
    write_tasks_md(generated_dir / "remediation-tasks.md", args.client, findings)

    write_report_md(
        generated_dir / "report-draft.md",
        args.client,
        run_dir,
        source_file,
        findings,
        args.project_filter,
    )

    write_report_docx(
        generated_dir / "report-draft.docx",
        args.client,
        run_dir,
        source_file,
        findings,
        args.project_filter,
        args.min_severity,
    )

    write_summary_json(generated_dir / "summary.json", findings)

    print(f"Parsed results: {source_file}")
    print(f"Minimum severity: {args.min_severity}")
    print(f"Findings found: {len(findings)}")
    print(f"Generated files: {generated_dir}")


if __name__ == "__main__":
    main()